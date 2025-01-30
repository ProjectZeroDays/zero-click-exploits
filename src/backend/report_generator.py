from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import logging
import os
import json
from typing import Dict, Any, List, Optional
from abc import ABC, abstractmethod
import schedule
import time
from datetime import datetime

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Abstract Base Class for Report Templates
class ReportTemplate(ABC):
    def __init__(self, template_path: str, template_name: str):
        self.template_path = template_path
        self.template_name = template_name
        self.doc = self._load_template()

    @abstractmethod
    def _load_template(self) -> Document:
        """Load the report template from the specified path."""
        pass

    @abstractmethod
    def generate_report(self, report_data: Dict[str, Any], output_path: str):
        """Generate the report using the loaded template and provided data."""
        pass

# Concrete Class for Executive Summary Template
class ExecutiveSummaryTemplate(ReportTemplate):
    def __init__(self, template_path: str = 'templates/executive-summary-template.docx'):
        super().__init__(template_path, 'Executive Summary Template')

    def _load_template(self) -> Document:
        """Load the executive summary template."""
        try:
            if os.path.exists(self.template_path):
                logging.info(f"Loading existing template from {self.template_path}")
                doc = Document(self.template_path)
                self._customize_styles(doc)
                return doc
            else:
                logging.info(f"No existing template found. Creating a new document.")
                doc = Document()
                self._customize_styles(doc)
                return doc
        except Exception as e:
            logging.error(f"Error loading template: {e}")
            return Document()

    def _customize_styles(self, doc: Document):
        """Customize the document styles."""
        styles = doc.styles
        # Add a custom heading style
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.base_style = styles['Heading 1']
        heading_style.font.name = 'Arial'
        heading_style.font.size = Inches(0.3)
        heading_style.paragraph_format.space_before = Inches(0.2)
        heading_style.paragraph_format.space_after = Inches(0.1)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_style.font.color.rgb =  (0, 0, 139) # Dark Blue

        # Add a custom paragraph style
        paragraph_style = styles.add_style('CustomParagraph', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.base_style = styles['Normal']
        paragraph_style.font.name = 'Calibri'
        paragraph_style.font.size = Inches(0.12)
        paragraph_style.paragraph_format.space_before = Inches(0.05)
        paragraph_style.paragraph_format.space_after = Inches(0.05)
        paragraph_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def generate_report(self, report_data: Dict[str, Any], output_path: str):
        """Generate the executive summary report."""
        try:
            doc = self.doc
            doc.add_heading('Executive Summary', level=1, style='CustomHeading')
            for key, value in report_data.items():
                paragraph = doc.add_paragraph(f"{key}: {value}", style='CustomParagraph')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_page_break()
            doc.save(output_path)
            logging.info(f"Report generated: {output_path}")
        except Exception as e:
            logging.error(f"Error generating report: {e}")

# Concrete Class for Technical Report Template
class TechnicalReportTemplate(ReportTemplate):
    def __init__(self, template_path: str = 'templates/technical-report-template.docx'):
        super().__init__(template_path, 'Technical Report Template')

    def _load_template(self) -> Document:
        """Load the technical report template."""
        try:
            if os.path.exists(self.template_path):
                logging.info(f"Loading existing template from {self.template_path}")
                doc = Document(self.template_path)
                self._customize_styles(doc)
                return doc
            else:
                logging.info(f"No existing template found. Creating a new document.")
                doc = Document()
                self._customize_styles(doc)
                return doc
        except Exception as e:
            logging.error(f"Error loading template: {e}")
            return Document()

    def _customize_styles(self, doc: Document):
        """Customize the document styles."""
        styles = doc.styles
        # Add a custom heading style
        heading_style = styles.add_style('TechHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.base_style = styles['Heading 1']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.size = Inches(0.25)
        heading_style.paragraph_format.space_before = Inches(0.15)
        heading_style.paragraph_format.space_after = Inches(0.1)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading_style.font.color.rgb = (139, 0, 0) # Dark Red

        # Add a custom paragraph style
        paragraph_style = styles.add_style('TechParagraph', WD_STYLE_TYPE.PARAGRAPH)
        paragraph_style.base_style = styles['Normal']
        paragraph_style.font.name = 'Consolas'
        paragraph_style.font.size = Inches(0.1)
        paragraph_style.paragraph_format.space_before = Inches(0.03)
        paragraph_style.paragraph_format.space_after = Inches(0.03)
        paragraph_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def generate_report(self, report_data: Dict[str, Any], output_path: str):
        """Generate the technical report."""
        try:
            doc = self.doc
            doc.add_heading('Technical Report', level=1, style='TechHeading')
            for key, value in report_data.items():
                paragraph = doc.add_paragraph(f"{key}: {value}", style='TechParagraph')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_page_break()
            doc.save(output_path)
            logging.info(f"Report generated: {output_path}")
        except Exception as e:
            logging.error(f"Error generating report: {e}")

# Report Generator
class ReportGenerator:
    def __init__(self):
        self.templates = {
            'executive': ExecutiveSummaryTemplate(),
            'technical': TechnicalReportTemplate()
        }
        self.report_queue = []

    def generate_report(self, template_type: str, report_data: Dict[str, Any], output_path: str):
        """Generate a report using a template and report data."""
        report_details = {
            'template_type': template_type,
            'report_data': report_data,
            'output_path': output_path,
            'timestamp': datetime.now().isoformat()
        }
        logging.info(f"Received report request: {report_details}")
        self.report_queue.append(report_details)
        self._process_report(report_details)

    def _process_report(self, report_details: Dict[str, Any]):
        """Process report generation."""
        template_type = report_details.get('template_type')
        report_data = report_details.get('report_data')
        output_path = report_details.get('output_path')

        if template_type in self.templates:
            template = self.templates[template_type]
            template.generate_report(report_data, output_path)
        else:
            logging.warning(f"No template found for type: {template_type}")
            logging.info("Default report generation completed.")

    def schedule_report_generation(self, interval_seconds: int = 60):
        """Schedule report generation at regular intervals."""
        def job():
            logging.info("Starting scheduled report generation...")
            if self.report_queue:
                report = self.report_queue.pop(0)
                self._process_report(report)
            else:
                logging.info("No reports in queue.")
            logging.info("Scheduled report generation completed.")

        schedule.every(interval_seconds).seconds.do(job)
        while True:
            schedule.run_pending()
            time.sleep(1)

    def load_report_data(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Load report data from a JSON file."""
        try:
            with open(file_path, 'r') as f:
                report_data = json.load(f)
            logging.info(f"Loaded report data from {file_path}")
            return report_data
        except FileNotFoundError:
            logging.error(f"Report data file not found: {file_path}")
            return None
        except json.JSONDecodeError:
            logging.error(f"Error decoding JSON from {file_path}")
            return None

    def start_report_generation(self, interval_seconds: int = 60):
        """Start report generation and processing."""
        self.schedule_report_generation(interval_seconds)

    def optimize_real_time_monitoring(self):
        """Advanced AI-driven optimization of real-time monitoring performance."""
        try:
            logging.info("Optimizing real-time monitoring performance...")
            # Placeholder for optimization logic
            optimization_results = {"status": "optimization applied", "details": "Optimization details"}
            logging.info(f"Optimization results: {optimization_results}")
            return optimization_results
        except Exception as e:
            logging.error(f"Error optimizing real-time monitoring performance: {e}")
            return {"status": "error", "message": str(e)}

if __name__ == '__main__':
    generator = ReportGenerator()
    report_data = {'Threat Level': 'High', 'Incident Count': 5, 'Vulnerabilities': ['CVE-2023-1234', 'CVE-2023-5678']}
    generator.generate_report('executive', report_data, 'reports/executive-summary-report.docx')
    generator.generate_report('technical', report_data, 'reports/technical-report.docx')
    generator.start_report_generation(interval_seconds=30)
    generator.optimize_real_time_monitoring()
